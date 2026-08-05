import os
from reportlab.pdfgen import canvas
from docx import Document

def create_pdf(filename, text, pages=1):
    c = canvas.Canvas(filename)
    for p in range(pages):
        c.drawString(100, 750, f"{text} - Page {p+1}")
        c.showPage()
    c.save()
    print(f"Created {filename}")

def create_docx(filename, text):
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == '__main__':
    os.makedirs('test_data', exist_ok=True)
    create_pdf('test_data/sample1.pdf', 'This is sample PDF 1')
    create_pdf('test_data/sample2.pdf', 'This is sample PDF 2')
    create_pdf('test_data/large.pdf', 'This is a large PDF for testing', pages=15)
    create_docx('test_data/sample.docx', 'This is a sample Word document for testing Word to PDF.')
